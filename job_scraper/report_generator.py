"""
report_generator.py — Generates .docx and .html daily job reports.
"""

import os
import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import CANDIDATE, OUTPUT

logger = logging.getLogger(__name__)

TODAY = datetime.utcnow().strftime("%Y-%m-%d")
NOW_PRETTY = datetime.utcnow().strftime("%B %d, %Y — %H:%M UTC")


# ─────────────────────────────────────────────────────────────
#  Shared helpers
# ─────────────────────────────────────────────────────────────

def score_badge(score) -> str:
    if score is None:
        return "N/A"
    s = int(score)
    if s >= 8:
        return f"⭐ {s}/10"
    if s >= 5:
        return f"🟡 {s}/10"
    return f"🔴 {s}/10"


def rec_badge(rec: str) -> str:
    return {"APPLY": "✅ APPLY", "CONSIDER": "🟡 CONSIDER", "SKIP": "❌ SKIP"}.get(rec, rec)


def _add_docx_divider(doc, color="2E75B6", size=6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p


# ─────────────────────────────────────────────────────────────
#  DOCX Report
# ─────────────────────────────────────────────────────────────

def generate_docx(professional: list[dict], general: list[dict], out_dir: str) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🔍 Daily Job Hunt Report")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{NOW_PRETTY}  |  {CANDIDATE['name']}  |  Aachen, NRW")
    r2.italic = True; r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    _add_docx_divider(doc)

    # ── Stats bar ──
    apply_count = sum(1 for j in professional + general if j.get("recommendation") == "APPLY")
    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        f"📊  Total found: {len(professional) + len(general)}  |  "
        f"Professional: {len(professional)}  |  General: {len(general)}  |  "
        f"Ready to apply: {apply_count}"
    )
    r3.bold = True; r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x1F, 0x6F, 0x3A)
    p3.paragraph_format.space_after = Pt(12)

    # ── Category 1 ──
    _docx_category_header(doc, "CATEGORY 1: Professional Part-Time Jobs", "1F497D")
    for job in professional:
        _docx_job_entry(doc, job)

    doc.add_page_break()

    # ── Category 2 ──
    _docx_category_header(doc, "CATEGORY 2: General Part-Time Jobs (NRW)", "1F6F3A")
    for job in general:
        _docx_job_entry(doc, job)

    # ── Footer ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated automatically by Job Hunter Bot powered by Claude (Anthropic) · {TODAY}")
    r.italic = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    path = os.path.join(out_dir, f"jobs_{TODAY}.docx")
    doc.save(path)
    logger.info(f"DOCX saved: {path}")
    return path


def _docx_category_header(doc, title: str, hex_color: str):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(15)
    rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    r.font.color.rgb = RGBColor(*rgb)
    _add_docx_divider(doc, color=hex_color)


def _docx_job_entry(doc, job: dict):
    score = job.get("relevance_score")
    rec = job.get("recommendation", "")

    # Title line
    p = doc.add_paragraph()
    r = p.add_run(f"{score_badge(score)}  {job.get('title', 'Unknown Title')}")
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

    # Meta line
    p2 = doc.add_paragraph()
    meta = (
        f"🏢 {job.get('company','?')}  |  📍 {job.get('location','?')}  |  "
        f"⏰ {job.get('employment_type', 'Part-time')}  |  "
        f"📌 {job.get('source','?')}  |  {rec_badge(rec)}"
    )
    r2 = p2.add_run(meta)
    r2.italic = True; r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2.paragraph_format.space_after = Pt(3)

    # Description
    desc = job.get("description", "")
    if desc:
        p3 = doc.add_paragraph()
        p3.add_run("Description: ").bold = True
        p3.runs[0].font.size = Pt(10)
        r3b = p3.add_run(desc[:400] + ("..." if len(desc) > 400 else ""))
        r3b.font.size = Pt(10)
        p3.paragraph_format.space_after = Pt(3)

    # Match reasons
    reasons = job.get("match_reasons", [])
    if reasons:
        p4 = doc.add_paragraph()
        rb = p4.add_run("Why You Match: ")
        rb.bold = True; rb.font.size = Pt(10)
        rb.font.color.rgb = RGBColor(0x1D, 0x6A, 0x39)
        r4b = p4.add_run(" · ".join(reasons))
        r4b.font.size = Pt(10)
        p4.paragraph_format.space_after = Pt(3)

    # Concerns
    concerns = job.get("concerns", [])
    if concerns:
        p5 = doc.add_paragraph()
        rc = p5.add_run("Concerns: ")
        rc.bold = True; rc.font.size = Pt(10)
        rc.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        r5b = p5.add_run(" · ".join(concerns))
        r5b.font.size = Pt(10)
        p5.paragraph_format.space_after = Pt(3)

    # Cover letter tip
    tip = job.get("cover_letter_tip", "")
    if tip:
        p6 = doc.add_paragraph()
        rt = p6.add_run("💡 Cover Letter Tip: ")
        rt.bold = True; rt.font.size = Pt(10)
        rt.font.color.rgb = RGBColor(0x7D, 0x3C, 0x98)
        r6b = p6.add_run(tip)
        r6b.font.size = Pt(10)
        p6.paragraph_format.space_after = Pt(3)

    # Apply URL
    url = job.get("apply_url", "")
    if url:
        p7 = doc.add_paragraph()
        p7.add_run("🔗 Apply: ").bold = True
        p7.runs[0].font.size = Pt(10)
        r7b = p7.add_run(url)
        r7b.font.size = Pt(10)
        r7b.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p7.paragraph_format.space_after = Pt(5)

    _add_docx_divider(doc, color="CCCCCC", size=2)


# ─────────────────────────────────────────────────────────────
#  HTML Report
# ─────────────────────────────────────────────────────────────

def generate_html(professional: list[dict], general: list[dict],
                  out_dir: str, docs_dir: str) -> str:
    apply_count = sum(1 for j in professional + general if j.get("recommendation") == "APPLY")
    total = len(professional) + len(general)

    jobs_pro_html = "".join(_html_job_card(j, "pro") for j in professional)
    jobs_gen_html = "".join(_html_job_card(j, "gen") for j in general)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>Job Hunter — {TODAY}</title>
  <style>
    :root {{
      --blue: #1F497D; --green: #1F6F3A; --purple: #7D3C98;
      --red: #C0392B; --gold: #D4AC0D; --light: #F4F6F9;
      --border: #DDE3ED; --text: #2C3E50; --muted: #7F8C8D;
    }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: 'Segoe UI', Arial, sans-serif; background: var(--light);
            color: var(--text); line-height: 1.6; }}
    header {{ background: var(--blue); color: #fff; padding: 28px 40px; }}
    header h1 {{ font-size: 1.8rem; margin-bottom: 4px; }}
    header p {{ opacity: .8; font-size: .9rem; }}
    .stats {{ display: flex; gap: 16px; padding: 20px 40px;
              background: #fff; border-bottom: 1px solid var(--border);
              flex-wrap: wrap; }}
    .stat-box {{ background: var(--light); border-radius: 8px; padding: 12px 20px;
                 text-align: center; min-width: 120px; }}
    .stat-box .num {{ font-size: 1.6rem; font-weight: 700; color: var(--blue); }}
    .stat-box .lbl {{ font-size: .75rem; color: var(--muted); text-transform: uppercase; }}
    .container {{ max-width: 1100px; margin: 0 auto; padding: 30px 20px; }}
    .section-title {{ font-size: 1.3rem; font-weight: 700; padding: 12px 0 4px;
                      border-bottom: 3px solid var(--blue); margin-bottom: 20px; }}
    .section-title.gen {{ border-color: var(--green); color: var(--green); }}
    .section-title.pro {{ color: var(--blue); }}
    .card {{ background: #fff; border: 1px solid var(--border); border-radius: 10px;
             padding: 20px 24px; margin-bottom: 16px;
             box-shadow: 0 2px 6px rgba(0,0,0,.05); transition: box-shadow .2s; }}
    .card:hover {{ box-shadow: 0 4px 14px rgba(0,0,0,.10); }}
    .card-header {{ display: flex; justify-content: space-between;
                    align-items: flex-start; flex-wrap: wrap; gap: 8px; }}
    .job-title {{ font-size: 1.05rem; font-weight: 700; color: var(--blue); }}
    .badges {{ display: flex; gap: 6px; flex-wrap: wrap; align-items: center; }}
    .badge {{ border-radius: 20px; padding: 2px 10px; font-size: .78rem;
              font-weight: 600; white-space: nowrap; }}
    .badge-apply  {{ background: #D5F5E3; color: #1A7A3A; }}
    .badge-consider {{ background: #FEF9E7; color: #9A7D0A; }}
    .badge-skip   {{ background: #FDEDEC; color: #A93226; }}
    .badge-score  {{ background: #EAF2FB; color: var(--blue); }}
    .badge-source {{ background: #F4ECF7; color: var(--purple); }}
    .meta {{ font-size: .83rem; color: var(--muted); margin: 6px 0 10px; }}
    .meta span {{ margin-right: 14px; }}
    .section-label {{ font-size: .78rem; font-weight: 700; text-transform: uppercase;
                      letter-spacing: .04em; margin-top: 10px; margin-bottom: 2px; }}
    .section-label.match {{ color: var(--green); }}
    .section-label.concern {{ color: var(--red); }}
    .section-label.tip {{ color: var(--purple); }}
    .section-label.desc {{ color: var(--text); }}
    .desc-text {{ font-size: .88rem; color: #444; }}
    .tag-list {{ display: flex; flex-wrap: wrap; gap: 6px; margin-top: 4px; }}
    .tag {{ background: var(--light); border: 1px solid var(--border);
            border-radius: 4px; padding: 1px 8px; font-size: .78rem; }}
    .apply-btn {{ display: inline-block; margin-top: 12px; background: var(--blue);
                  color: #fff; padding: 7px 18px; border-radius: 6px;
                  text-decoration: none; font-size: .85rem; font-weight: 600; }}
    .apply-btn:hover {{ background: #16335A; }}
    .apply-btn.gen {{ background: var(--green); }}
    .apply-btn.gen:hover {{ background: #155228; }}
    footer {{ text-align: center; padding: 30px; color: var(--muted); font-size: .82rem;
              border-top: 1px solid var(--border); background: #fff; margin-top: 20px; }}
    .filter-bar {{ display: flex; gap: 10px; margin-bottom: 20px; flex-wrap: wrap; }}
    .filter-btn {{ border: 1px solid var(--border); background: #fff; border-radius: 20px;
                   padding: 5px 14px; cursor: pointer; font-size: .83rem; transition: all .2s; }}
    .filter-btn.active, .filter-btn:hover {{ background: var(--blue); color: #fff;
                                              border-color: var(--blue); }}
    @media(max-width:600px) {{ header {{ padding: 16px 18px; }}
      .stats {{ padding: 12px 18px; }} .container {{ padding: 16px 12px; }} }}
  </style>
</head>
<body>

<header>
  <h1>🔍 Daily Job Hunt Report</h1>
  <p>{NOW_PRETTY} &nbsp;·&nbsp; {CANDIDATE['name']} &nbsp;·&nbsp; Aachen, NRW, Germany</p>
</header>

<div class="stats">
  <div class="stat-box"><div class="num">{total}</div><div class="lbl">Total Jobs</div></div>
  <div class="stat-box"><div class="num">{len(professional)}</div><div class="lbl">Professional</div></div>
  <div class="stat-box"><div class="num">{len(general)}</div><div class="lbl">General</div></div>
  <div class="stat-box" style="border-left:4px solid #1A7A3A">
    <div class="num" style="color:#1A7A3A">{apply_count}</div>
    <div class="lbl">Apply Now</div>
  </div>
</div>

<div class="container">

  <!-- Filter bar -->
  <div class="filter-bar">
    <button class="filter-btn active" onclick="filterJobs('all')">All Jobs</button>
    <button class="filter-btn" onclick="filterJobs('APPLY')">✅ Apply</button>
    <button class="filter-btn" onclick="filterJobs('CONSIDER')">🟡 Consider</button>
    <button class="filter-btn" onclick="filterJobs('pro')">Professional</button>
    <button class="filter-btn" onclick="filterJobs('gen')">General</button>
  </div>

  <div class="section-title pro">📋 Category 1 — Professional Part-Time Jobs</div>
  <div id="pro-jobs">{jobs_pro_html}</div>

  <div class="section-title gen" style="margin-top:30px">📦 Category 2 — General Part-Time Jobs</div>
  <div id="gen-jobs">{jobs_gen_html}</div>

</div>

<footer>
  Generated automatically every day at 12:00 PM UTC by
  <strong>Job Hunter Bot</strong> powered by <strong>Claude (Anthropic)</strong>
  &nbsp;·&nbsp; Report date: {TODAY}
  &nbsp;·&nbsp;
  <a href="jobs_{TODAY}.docx" style="color:var(--blue)">⬇ Download Word (.docx)</a>
</footer>

<script>
function filterJobs(filter) {{
  document.querySelectorAll('.filter-btn').forEach(b => b.classList.remove('active'));
  event.target.classList.add('active');
  document.querySelectorAll('.card').forEach(card => {{
    if (filter === 'all') {{ card.style.display = ''; return; }}
    if (filter === 'pro') {{ card.style.display = card.dataset.cat === 'pro' ? '' : 'none'; return; }}
    if (filter === 'gen') {{ card.style.display = card.dataset.cat === 'gen' ? '' : 'none'; return; }}
    card.style.display = card.dataset.rec === filter ? '' : 'none';
  }});
}}
</script>
</body>
</html>"""

    # Save to docs/ (GitHub Pages) and output/reports/
    for directory in [out_dir, docs_dir]:
        Path(directory).mkdir(parents=True, exist_ok=True)

    html_filename = f"jobs_{TODAY}.html"
    report_path = os.path.join(out_dir, html_filename)
    docs_path = os.path.join(docs_dir, "index.html")   # GitHub Pages entry point

    with open(report_path, "w", encoding="utf-8") as f:
        f.write(html)
    with open(docs_path, "w", encoding="utf-8") as f:
        f.write(html)

    logger.info(f"HTML saved: {report_path}")
    logger.info(f"GitHub Pages index updated: {docs_path}")
    return report_path


def _html_job_card(job: dict, cat: str) -> str:
    rec = job.get("recommendation", "CONSIDER")
    score = job.get("relevance_score", "N/A")
    rec_class = {"APPLY": "badge-apply", "CONSIDER": "badge-consider", "SKIP": "badge-skip"}.get(rec, "badge-consider")
    score_color = "#1A7A3A" if (isinstance(score, int) and score >= 8) else ("#9A7D0A" if isinstance(score, int) and score >= 5 else "#A93226")
    score_label = f"{score}/10" if isinstance(score, int) else "N/A"

    reasons_html = ""
    if job.get("match_reasons"):
        items = "".join(f'<span class="tag">✓ {r}</span>' for r in job["match_reasons"])
        reasons_html = f'<div class="section-label match">Why You Match</div><div class="tag-list">{items}</div>'

    concerns_html = ""
    if job.get("concerns"):
        items = "".join(f'<span class="tag">⚠ {c}</span>' for c in job["concerns"])
        concerns_html = f'<div class="section-label concern">Concerns</div><div class="tag-list">{items}</div>'

    skills_html = ""
    if job.get("key_skills_required"):
        items = "".join(f'<span class="tag">{s}</span>' for s in job["key_skills_required"])
        skills_html = f'<div class="section-label desc">Skills Required</div><div class="tag-list">{items}</div>'

    tip_html = ""
    if job.get("cover_letter_tip"):
        tip_html = f'<div class="section-label tip">💡 Cover Letter Tip</div><p class="desc-text">{job["cover_letter_tip"]}</p>'

    desc = (job.get("description") or "")[:350]
    if len(job.get("description", "")) > 350:
        desc += "..."

    apply_url = job.get("apply_url", "#")
    btn_class = "apply-btn gen" if cat == "gen" else "apply-btn"

    german_flag = "🇩🇪 German required" if job.get("german_required") else ""

    return f"""
<div class="card" data-cat="{cat}" data-rec="{rec}">
  <div class="card-header">
    <div class="job-title">{job.get('title', 'Unknown')}</div>
    <div class="badges">
      <span class="badge badge-score" style="color:{score_color}">⭐ {score_label}</span>
      <span class="badge {rec_class}">{rec_badge(rec)}</span>
      <span class="badge badge-source">{job.get('source','?')}</span>
      {f'<span class="badge badge-consider">{german_flag}</span>' if german_flag else ''}
    </div>
  </div>
  <div class="meta">
    <span>🏢 {job.get('company','?')}</span>
    <span>📍 {job.get('location','?')}</span>
    <span>⏰ {job.get('employment_type','Part-time')}</span>
    <span>📅 {job.get('posted','')}</span>
  </div>
  <div class="section-label desc">Description</div>
  <p class="desc-text">{desc}</p>
  {reasons_html}
  {concerns_html}
  {skills_html}
  {tip_html}
  <a href="{apply_url}" target="_blank" class="{btn_class}">🔗 Apply Now</a>
</div>"""


def rec_badge(rec: str) -> str:
    return {"APPLY": "✅ APPLY", "CONSIDER": "🟡 CONSIDER", "SKIP": "❌ SKIP"}.get(rec, rec)


# ─────────────────────────────────────────────────────────────
#  JSON dump (for future use / data analysis)
# ─────────────────────────────────────────────────────────────

def save_json(professional: list[dict], general: list[dict], out_dir: str) -> str:
    data = {
        "generated_at": datetime.utcnow().isoformat(),
        "candidate": CANDIDATE["name"],
        "professional_jobs": professional,
        "general_jobs": general,
    }
    path = os.path.join(out_dir, f"jobs_{TODAY}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.info(f"JSON saved: {path}")
    return path
